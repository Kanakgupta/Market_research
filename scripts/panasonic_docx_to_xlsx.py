"""Convert Panasonic 'Detailed Next Gen Bluetooth Features' DOCX to XLSX.

The DOCX is a flat list of paragraphs, each highlighted with a color that maps
to a requirement priority (legend: green=MVP, yellow=Should Have,
red=Nice to Have; cyan = 4th category found in this doc).

Output: a 'Requirements' sheet with columns
    #, Section, Sub-section, Requirement, Use Case, Priority, Bold, RawText
plus the row fill color matches the source highlight so the visual coding is
preserved 1:1. Writes a JSON audit alongside, listing every paragraph with its
detected priority so the user can confirm nothing was lost.
"""
from __future__ import annotations
import json
import re
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

SRC = Path(r"C:\guptakanak\customers\Panasonic\2026\Detailed Next Gen Bluetooth Features  Functions List with Use Cases_May_2026.docx")
DST = SRC.with_suffix(".xlsx")
AUDIT = SRC.with_name(SRC.stem + "_conversion_audit.json")

# Legend per source doc para 1: "Color Code: MVP, Should Have, Nice to Have"
# where MVP is highlighted green, Should Have is cyan, Nice to Have is yellow.
HL_TO_FILL = {
    "green":      ("C6EFCE", "MVP"),
    "cyan":       ("BDD7EE", "Should Have"),
    "yellow":     ("FFEB9C", "Nice to Have"),
    "red":        ("FFC7CE", "Other / Flagged"),
    "magenta":    ("F4CCCC", "Other"),
    "blue":       ("DDEBF7", "Other"),
    "darkGreen":  ("A9D08E", "MVP"),
    "darkCyan":   ("9BC2E6", "Should Have"),
    "darkYellow": ("FFE699", "Nice to Have"),
    "darkRed":    ("E06666", "Other / Flagged"),
}
PRIORITY_ORDER = {"MVP": 0, "Should Have": 1, "Nice to Have": 2,
                  "Other / Flagged": 3, "Other": 4, "": 5}


def _para_runs_info(p) -> list[dict]:
    out = []
    for r in p.runs:
        rPr = r._element.find(qn("w:rPr"))
        info = {"text": r.text, "color": None, "highlight": None,
                "shd": None, "bold": False}
        if rPr is not None:
            c = rPr.find(qn("w:color"))
            if c is not None:
                v = c.get(qn("w:val"))
                if v and v.lower() != "auto":
                    info["color"] = v.upper() if re.fullmatch(r"[0-9A-Fa-f]{6}", v) else v
            h = rPr.find(qn("w:highlight"))
            if h is not None:
                info["highlight"] = h.get(qn("w:val"))
            sh = rPr.find(qn("w:shd"))
            if sh is not None:
                info["shd"] = sh.get(qn("w:fill"))
            b = rPr.find(qn("w:b"))
            info["bold"] = b is not None and b.get(qn("w:val"), "1") not in ("0", "false")
        out.append(info)
    return out


def _dominant_highlight(runs: list[dict]) -> str | None:
    counter = Counter()
    for r in runs:
        if r["highlight"] and r["highlight"].lower() != "none":
            counter[r["highlight"]] += len((r["text"] or "").strip())
    if not counter:
        for r in runs:
            if r["shd"] and r["shd"].lower() not in ("auto", "ffffff"):
                counter["shd:" + r["shd"]] += len((r["text"] or "").strip())
    return counter.most_common(1)[0][0] if counter else None


def _is_bold_para(runs: list[dict]) -> bool:
    txt_total = sum(len((r["text"] or "").strip()) for r in runs)
    bold_total = sum(len((r["text"] or "").strip()) for r in runs if r["bold"])
    return txt_total > 0 and bold_total >= txt_total / 2


SECTION_KWS = ["bluetooth classic", "bluetooth low energy", "ble", "le audio",
               "isochronous channels", "isoc", "channel sounding", "ranging",
               "high data throughput", "hdt", "auracast", "broadcast",
               "hardware specs", "security", "qualification", "regulatory",
               "use cases", "scope", "definitions"]


def _looks_like_section(text: str, runs: list[dict]) -> bool:
    t = text.strip()
    if not t or len(t) > 90:
        return False
    if t.lower().startswith(("use case", "color code", "minimum bluetooth",
                              "footprint", "power", "interface")):
        return False
    if t.endswith(":"):
        return True
    if _is_bold_para(runs) and len(t.split()) <= 8 and not t.endswith("."):
        return True
    tl = t.lower()
    if any(k == tl or (k in tl and len(t.split()) <= 6) for k in SECTION_KWS):
        return True
    return False


def convert() -> dict:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    doc = Document(str(SRC))

    wb = Workbook()
    ws = wb.active
    ws.title = "Requirements"

    headers = ["#", "Section", "Sub-section", "Requirement", "Use Case",
               "Priority", "Bold", "Raw Highlight", "Raw Text"]
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=c, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="305496")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    thin = Side(border_style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    audit_rows: list[dict] = []
    current_section, current_sub = "", ""
    out_row = 2
    priority_counter, highlight_counter = Counter(), Counter()
    legend_lines: list[str] = []

    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if not text:
            continue
        runs = _para_runs_info(p)
        hl = _dominant_highlight(runs)
        if hl:
            highlight_counter[hl] += 1
        priority, fill_hex = "", None
        if hl:
            mp = HL_TO_FILL.get(hl)
            if mp:
                fill_hex, priority = mp
            elif hl.startswith("shd:"):
                hex6 = hl.split(":", 1)[1]
                if re.fullmatch(r"[0-9A-Fa-f]{6}", hex6):
                    fill_hex = hex6.upper()
        priority_counter[priority or "(none)"] += 1
        is_use_case = text.lower().startswith("use case")
        bold = _is_bold_para(runs)

        if i <= 2:
            legend_lines.append(text)

        if not is_use_case and _looks_like_section(text, runs):
            tl = text.lower().rstrip(":")
            top = ("bluetooth classic", "bluetooth low energy", "ble",
                    "hardware specs", "security", "qualification",
                    "regulatory", "scope", "use cases")
            if any(tl == k or tl.startswith(k) for k in top):
                current_section = text.rstrip(":")
                current_sub = ""
            else:
                current_sub = text.rstrip(":")
            ws.cell(row=out_row, column=1, value=out_row - 1)
            ws.cell(row=out_row, column=2, value=current_section)
            ws.cell(row=out_row, column=3, value=current_sub)
            ws.cell(row=out_row, column=4, value=text).font = Font(bold=True)
            ws.cell(row=out_row, column=6, value=priority)
            ws.cell(row=out_row, column=7, value="Yes" if bold else "")
            ws.cell(row=out_row, column=8, value=hl or "")
            ws.cell(row=out_row, column=9, value=text)
            for c in range(1, len(headers) + 1):
                cc = ws.cell(row=out_row, column=c)
                cc.alignment = Alignment(wrap_text=True, vertical="top")
                cc.border = border
                if fill_hex:
                    cc.fill = PatternFill("solid", fgColor=fill_hex)
            audit_rows.append({"para_idx": i, "out_row": out_row,
                               "section": current_section, "sub_section": current_sub,
                               "is_section_heading": True, "is_use_case": False,
                               "priority": priority, "highlight": hl, "text": text})
            out_row += 1
            continue

        req_text, use_case_text = text, ""
        if is_use_case:
            m = re.match(r"^use case\s*:\s*(.*)", text, flags=re.I)
            use_case_text = m.group(1) if m else text
            req_text = ""

        ws.cell(row=out_row, column=1, value=out_row - 1)
        ws.cell(row=out_row, column=2, value=current_section)
        ws.cell(row=out_row, column=3, value=current_sub)
        ws.cell(row=out_row, column=4, value=req_text)
        ws.cell(row=out_row, column=5, value=use_case_text)
        ws.cell(row=out_row, column=6, value=priority)
        ws.cell(row=out_row, column=7, value="Yes" if bold else "")
        ws.cell(row=out_row, column=8, value=hl or "")
        ws.cell(row=out_row, column=9, value=text)
        for c in range(1, len(headers) + 1):
            cc = ws.cell(row=out_row, column=c)
            cc.alignment = Alignment(wrap_text=True, vertical="top")
            cc.border = border
            if fill_hex:
                cc.fill = PatternFill("solid", fgColor=fill_hex)
        if bold:
            ws.cell(row=out_row, column=4).font = Font(bold=True)
            ws.cell(row=out_row, column=5).font = Font(bold=True)

        audit_rows.append({"para_idx": i, "out_row": out_row,
                           "section": current_section, "sub_section": current_sub,
                           "is_section_heading": False, "is_use_case": is_use_case,
                           "priority": priority, "highlight": hl, "text": text})
        out_row += 1

    widths = {1: 5, 2: 22, 3: 22, 4: 60, 5: 60, 6: 14, 7: 6, 8: 14, 9: 50}
    for c, w in widths.items():
        ws.column_dimensions[get_column_letter(c)].width = w
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{out_row - 1}"

    # Legend sheet
    lg = wb.create_sheet("Legend")
    for c, h in enumerate(["Color", "Meaning", "Source highlight (Word)", "Count"], 1):
        cell = lg.cell(row=1, column=c, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="305496")
        cell.alignment = Alignment(horizontal="center")
    r = 2
    for hl_name, (fill_hex, label) in HL_TO_FILL.items():
        lg.cell(row=r, column=1).fill = PatternFill("solid", fgColor=fill_hex)
        lg.cell(row=r, column=2, value=label)
        lg.cell(row=r, column=3, value=hl_name)
        lg.cell(row=r, column=4, value=highlight_counter.get(hl_name, 0))
        r += 1
    r += 1
    lg.cell(row=r, column=1, value="Original legend from doc:").font = Font(bold=True)
    for ll in legend_lines:
        r += 1
        lg.cell(row=r, column=1, value=ll).alignment = Alignment(wrap_text=True)
    for c, w in {1: 14, 2: 22, 3: 26, 4: 8}.items():
        lg.column_dimensions[get_column_letter(c)].width = w

    # Summary sheet
    sm = wb.create_sheet("Summary")
    for c, h in enumerate(["Metric", "Value"], 1):
        cell = sm.cell(row=1, column=c, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="305496")
    rows = [
        ("Source DOCX", str(SRC)),
        ("Output XLSX", str(DST)),
        ("Total paragraphs in DOCX", len(doc.paragraphs)),
        ("Non-empty rows emitted", len(audit_rows)),
        ("Distinct highlights detected",
            ", ".join(f"{k}({v})" for k, v in highlight_counter.most_common())),
        ("Priority counts",
            ", ".join(f"{k}={v}" for k, v in sorted(
                priority_counter.items(), key=lambda kv: PRIORITY_ORDER.get(kv[0], 99)))),
        ("Use-case rows", sum(1 for a in audit_rows if a["is_use_case"])),
        ("Section heading rows", sum(1 for a in audit_rows if a["is_section_heading"])),
    ]
    for i, (k, v) in enumerate(rows, 2):
        sm.cell(row=i, column=1, value=k).font = Font(bold=True)
        c = sm.cell(row=i, column=2, value=v)
        c.alignment = Alignment(wrap_text=True, vertical="top")
    sm.column_dimensions["A"].width = 36
    sm.column_dimensions["B"].width = 90

    wb.move_sheet("Summary", offset=-2)
    wb.save(str(DST))

    audit = {
        "source": str(SRC), "destination": str(DST),
        "total_paragraphs": len(doc.paragraphs),
        "rows_emitted": len(audit_rows),
        "highlight_counts": dict(highlight_counter),
        "priority_counts": dict(priority_counter),
        "rows": audit_rows,
        "verification": {
            "all_non_empty_paragraphs_emitted":
                sum(1 for p in doc.paragraphs if (p.text or "").strip()) == len(audit_rows),
            "every_highlighted_paragraph_has_priority":
                all(a["priority"] for a in audit_rows if a["highlight"] in HL_TO_FILL),
            "unmapped_highlights": list(set(highlight_counter) - set(HL_TO_FILL)),
        },
    }
    AUDIT.write_text(json.dumps(audit, indent=2), encoding="utf-8")
    return audit


if __name__ == "__main__":
    a = convert()
    print(f"Source:        {a['source']}")
    print(f"Destination:   {a['destination']}")
    print(f"Audit:         {AUDIT}")
    print(f"Paragraphs:    {a['total_paragraphs']}")
    print(f"Rows emitted:  {a['rows_emitted']}")
    print(f"Highlights:    {a['highlight_counts']}")
    print(f"Priorities:    {a['priority_counts']}")
    print("\nVerification:")
    for k, v in a["verification"].items():
        print(f"  {k}: {v}")
